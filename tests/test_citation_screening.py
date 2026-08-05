import io
import unittest
from unittest.mock import patch

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from citation_screening.parsers import parse_manuscript
from citation_screening.fulltext_review import execute_fulltext_review, prepare_fulltext_review
from citation_screening.author_report import execute_author_report, prepare_author_report
from citation_screening.pipeline import execute_screening, prepare_screening, run_screening
from citation_screening.reports import build_html_report
from citation_screening.services.deepseek import _extract_json


class ParserTests(unittest.TestCase):
    @staticmethod
    def _append_superscript_run(parent, text):
        run = OxmlElement("w:r")
        props = OxmlElement("w:rPr")
        align = OxmlElement("w:vertAlign")
        align.set(qn("w:val"), "superscript")
        props.append(align)
        run.append(props)
        value = OxmlElement("w:t")
        value.text = text
        run.append(value)
        parent.append(run)

    def test_word_superscript_range_and_reference_section(self):
        document = Document()
        paragraph = document.add_paragraph()
        paragraph.add_run("Drug A reduced inflammation ")
        first = paragraph.add_run("1")
        first.font.superscript = True
        dash = paragraph.add_run("–")
        dash.font.superscript = True
        last = paragraph.add_run("2")
        last.font.superscript = True
        paragraph.add_run(".")
        document.add_paragraph("References")
        document.add_paragraph("[1] This bibliography entry must not become a citation sentence.")
        buffer = io.BytesIO()
        document.save(buffer)

        sentences, refs = parse_manuscript(
            buffer.getvalue(), "paper.docx", "1. First ref\n2. Second ref"
        )
        self.assertEqual(len(sentences), 1)
        self.assertEqual(sentences[0]["citations"], [{"rid": "B1"}, {"rid": "B2"}])
        self.assertEqual(set(refs), {"B1", "B2"})

    def test_word_context_window(self):
        document = Document()
        document.add_paragraph("Earlier work established the mechanism. Drug A reduced inflammation [1]. The effect depended on dose.")
        buffer = io.BytesIO()
        document.save(buffer)
        sentences, _ = parse_manuscript(buffer.getvalue(), "paper.docx", "1. First ref")
        self.assertEqual(sentences[0]["context_before"], "Earlier work established the mechanism.")
        self.assertEqual(sentences[0]["sentence_text"], "Drug A reduced inflammation [1].")
        self.assertEqual(sentences[0]["context_after"], "The effect depended on dose.")

    def test_word_field_and_content_control_citations(self):
        document = Document()
        field_paragraph = document.add_paragraph("Field citation ")
        begin_run = OxmlElement("w:r")
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        begin_run.append(begin)
        field_paragraph._p.append(begin_run)
        instruction_run = OxmlElement("w:r")
        instruction = OxmlElement("w:instrText")
        instruction.text = " ADDIN EN.CITE DATA "
        instruction_run.append(instruction)
        field_paragraph._p.append(instruction_run)
        separate_run = OxmlElement("w:r")
        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        separate_run.append(separate)
        field_paragraph._p.append(separate_run)
        self._append_superscript_run(field_paragraph._p, "1")
        end_run = OxmlElement("w:r")
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        end_run.append(end)
        field_paragraph._p.append(end_run)
        field_paragraph.add_run(".")

        control_paragraph = document.add_paragraph("Control citation ")
        content_control = OxmlElement("w:sdt")
        content = OxmlElement("w:sdtContent")
        self._append_superscript_run(content, "2")
        content_control.append(content)
        control_paragraph._p.append(content_control)
        control_paragraph.add_run(".")

        buffer = io.BytesIO()
        document.save(buffer)
        sentences, _ = parse_manuscript(buffer.getvalue(), "fields.docx", "1. First\n2. Second")
        self.assertEqual(len(sentences), 2)
        self.assertEqual(sentences[0]["citations"], [{"rid": "B1"}])
        self.assertEqual(sentences[1]["citations"], [{"rid": "B2"}])
        self.assertNotIn("ADDIN EN.CITE", sentences[0]["sentence_text"])

    def test_namespaced_nlm_xml(self):
        xml = b'''<article xmlns="http://jats.nlm.nih.gov"><body><p>Claim <xref ref-type="bibr" rid="B1">[1]</xref>.</p></body><back><ref-list><ref id="B1"><label>1</label><element-citation><article-title>Article title</article-title><pub-id pub-id-type="doi">10.1/test</pub-id></element-citation></ref></ref-list></back></article>'''
        sentences, refs = parse_manuscript(xml, "paper.xml")
        self.assertEqual(sentences[0]["citations"], [{"rid": "B1"}])
        self.assertEqual(refs["B1"]["title"], "Article title")
        self.assertEqual(refs["B1"]["doi"], "10.1/test")


class PipelineTests(unittest.TestCase):
    def test_unrelated_header_content_category_and_percentage(self):
        parsed = _extract_json('{"result":"引用无关内容","reason":"作者单位编号"}')
        self.assertEqual(parsed["result"], "引用无关内容")
        html = build_html_report("paper.docx", [
            {"result": "匹配", "sentence_text": "Scientific claim.", "label": "1"},
            {"result": "引用无关内容", "sentence_text": "[1] University Hospital", "label": "1"},
        ])
        self.assertIn("匹配 <span class=\"percent\">100.0%</span>", html)
        self.assertIn("占比统计分母：1 条", html)
        self.assertIn('<option value="引用无关内容">', html)

    @patch("citation_screening.pipeline.fetch_metadata")
    def test_missing_abstract_is_separate_from_doubt(self, fetch_metadata):
        fetch_metadata.return_value = {
            "title": "Article title", "abstract": "", "authors": [],
            "metadata_source": "Crossref", "metadata_error": "",
        }
        xml = b'''<article><body><p>Claim <xref ref-type="bibr" rid="B1">[1]</xref>.</p></body><back><ref-list><ref id="B1"><label>1</label><article-title>Article title</article-title></ref></ref-list></back></article>'''
        data = run_screening(xml, "paper.xml")
        self.assertEqual(data["statistics"]["未获取数据"], 1)
        self.assertEqual(data["statistics"]["存疑"], 0)
        self.assertEqual(data["results"][0]["result"], "未获取数据")
        self.assertIn('<option value="未获取数据">', data["html"])

    @patch("citation_screening.pipeline.screen_pair")
    @patch("citation_screening.pipeline.fetch_metadata")
    def test_pipeline_and_html_without_external_calls(self, fetch_metadata, screen_pair):
        fetch_metadata.return_value = {
            "title": "Article title", "abstract": "Supporting abstract",
            "authors": [], "metadata_source": "PubMed", "metadata_error": "",
        }
        screen_pair.return_value = {"result": "匹配", "reason": "摘要支持主要主张。"}
        xml = b'''<article><body><p>Claim <xref ref-type="bibr" rid="B1">[1]</xref>.</p></body><back><ref-list><ref id="B1"><label>1</label><element-citation><article-title>Article title</article-title></element-citation></ref></ref-list></back></article>'''
        data = run_screening(xml, "paper.xml")
        self.assertEqual(data["statistics"]["匹配"], 1)
        self.assertIn("REF-C 引用内容初筛报告", data["html"])
        self.assertIn('id="statusFilter"', data["html"])
        self.assertIn('id="sourceFilter"', data["html"])
        self.assertIn("目标引用句", data["html"])
        self.assertEqual(data["results"][0]["result"], "匹配")

    @patch("citation_screening.pipeline.screen_pair")
    @patch("citation_screening.pipeline.fetch_metadata")
    def test_prepare_never_calls_ai_and_execute_uses_estimate(self, fetch_metadata, screen_pair):
        fetch_metadata.return_value = {
            "title": "Article title", "abstract": "Supporting abstract",
            "authors": [], "metadata_source": "PubMed", "metadata_error": "",
        }
        screen_pair.return_value = {
            "result": "匹配", "reason": "摘要支持。", "api_called": True,
        }
        xml = b'''<article><body><p>Before. Claim <xref ref-type="bibr" rid="B1">[1]</xref>. After.</p></body><back><ref-list><ref id="B1"><label>1</label><article-title>Article title</article-title></ref></ref-list></back></article>'''
        prepared = prepare_screening(xml, "paper.xml")
        screen_pair.assert_not_called()
        self.assertEqual(prepared["estimated_calls"], 1)
        result = execute_screening(prepared)
        self.assertEqual(result["actual_calls"], 1)
        self.assertEqual(screen_pair.call_args.args[2], "Before.")
        self.assertEqual(screen_pair.call_args.args[3], "After.")


class FulltextReviewTests(unittest.TestCase):
    @patch("citation_screening.fulltext_review.check_fulltext_paragraph")
    @patch("citation_screening.fulltext_review.fetch_open_fulltext")
    def test_doubtful_only_and_stop_after_support(self, fetch_fulltext, check_paragraph):
        fetch_fulltext.return_value = {
            "source": "PMC XML",
            "source_url": "https://example.test/fulltext",
            "paragraphs": [
                {"section": "Methods", "text": "Unrelated methods paragraph " * 8, "page": None},
                {"section": "Results", "text": "Drug A significantly reduced inflammation in the treated group. " * 3, "page": None},
                {"section": "Discussion", "text": "A later paragraph should never be checked. " * 4, "page": None},
            ],
        }
        check_paragraph.side_effect = [
            {"decision": "不支持", "reason": "没有直接证据", "api_called": True},
            {"decision": "支持", "reason": "结果段直接报告该效应", "api_called": True},
        ]
        base = {
            "filename": "paper.xml",
            "results": [
                {
                    "rid": "B1", "label": "1", "sentence_text": "Drug A reduced inflammation.",
                    "title": "Drug A study", "result": "存疑", "reason": "摘要证据不足",
                },
                {
                    "rid": "B2", "label": "2", "sentence_text": "A known claim.",
                    "title": "Known study", "result": "匹配", "reason": "摘要支持",
                },
            ],
        }
        prepared = prepare_fulltext_review(base, max_paragraphs=3, max_workers=1)
        self.assertEqual(prepared["doubtful_count"], 1)
        self.assertEqual(prepared["fulltexts_found"], 1)
        self.assertEqual(prepared["estimated_calls"], 3)

        # Execution preserves this ranked order and must stop before candidate 3.
        prepared["review_items"][0]["candidates"] = [
            {"section": "Methods", "text": "First candidate has no direct evidence.", "page": 1},
            {"section": "Results", "text": "Drug A significantly reduced inflammation in the treated group.", "page": 2},
            {"section": "Discussion", "text": "This candidate must not be sent.", "page": 3},
        ]

        reviewed = execute_fulltext_review(prepared, max_workers=1)
        self.assertEqual(check_paragraph.call_count, 2)
        self.assertEqual(reviewed["actual_calls"], 2)
        self.assertEqual(reviewed["results"][0]["result"], "匹配")
        self.assertIn("Drug A significantly", reviewed["results"][0]["evidence_text"])
        self.assertEqual(reviewed["results"][1]["result"], "匹配")
        self.assertIn("全文复核", reviewed["html"])
        self.assertIn("已全文复核", reviewed["html"])


class AuthorReportTests(unittest.TestCase):
    @patch("citation_screening.author_report._generate_batch")
    def test_only_doubt_and_mismatch_are_reported_with_combined_context(self, generate_batch):
        data = {"filename": "paper.docx", "results": [
            {"result": "匹配", "label": "1", "sentence_text": "Supported."},
            {"result": "存疑", "label": "2", "context_before": "Before.",
             "sentence_text": "Target [2].", "context_after": "After.",
             "reason": "摘要证据不足", "title": "Paper two"},
            {"result": "领域不符", "label": "3", "sentence_text": "Another claim [3].",
             "reason": "研究对象不同", "title": "Paper three"},
            {"result": "未获取数据", "label": "4", "sentence_text": "Missing [4]."},
            {"result": "引用无关内容", "label": "5", "sentence_text": "University [5]."},
        ]}
        prepared = prepare_author_report(data)
        self.assertEqual(len(prepared["items"]), 2)
        self.assertEqual(prepared["estimated_calls"], 1)
        self.assertEqual(prepared["items"][0]["relevant_text"], "Before. Target [2]. After.")
        generate_batch.return_value = {
            "summary": "Two references require the authors' attention.",
            "items": [
                {"item_id": "1", "concern": "Support is unclear.", "suggested_action": "Please verify the reference."},
                {"item_id": "2", "concern": "The population differs.", "suggested_action": "Please cite a more suitable study."},
            ],
        }
        report = execute_author_report(prepared)
        self.assertEqual(report["actual_calls"], 1)
        self.assertIn("Reference Check Report", report["html"])
        self.assertIn("Reference 2", report["html"])
        self.assertIn("Before. Target [2]. After.", report["html"])
        self.assertNotIn("University [5].", report["html"])


if __name__ == "__main__":
    unittest.main()
